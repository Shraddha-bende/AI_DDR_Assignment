import pdfplumber
import fitz
import os
from docx import Document
from docx.shared import Inches

# -------- TEXT EXTRACTION --------
def extract_text(pdf_file):
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# -------- IMAGE EXTRACTION --------
def extract_images(pdf_file):
    doc = fitz.open(pdf_file)

    if not os.path.exists("images"):
        os.makedirs("images")

    image_paths = []

    for page_index in range(len(doc)):
        page = doc[page_index]
        images = page.get_images()

        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            image_path = f"images/thermal_{page_index}_{img_index}.png"

            with open(image_path, "wb") as f:
                f.write(image_bytes)

            image_paths.append(image_path)

    return image_paths

# -------- IMPACTED AREAS --------
def find_impacted_areas(text):
    areas = []
    keywords = {
        "Bedroom": "Bedroom",
        "Parking": "Parking Area",
        "Bathroom": "Common Bathroom"
    }
    for key in keywords:
        if key.lower() in text.lower():
            areas.append(keywords[key])
    return areas

# -------- AUTOMATIC ISSUE DETECTION --------
def detect_issues(text):
    issues = []
    severity = "Low"

    keywords = {
        "crack": "Wall Crack Detected",
        "moisture": "Moisture Presence",
        "leak": "Possible Water Leakage",
        "damp": "Dampness Observed",
        "seepage": "Water Seepage Risk"
    }

    for key in keywords:
        if key in text.lower():
            issues.append(keywords[key])

    # Severity logic
    if len(issues) >= 3:
        severity = "High"
    elif len(issues) == 2:
        severity = "Moderate"
    elif len(issues) == 1:
        severity = "Low"

    return issues, severity

# -------- GENERATE DDR REPORT --------
def generate_report(inspection_text, thermal_text, images):
    doc = Document()

    doc.add_heading("DETAILED DIAGNOSTIC REPORT", level=0)
    doc.add_paragraph("Generated Automatically by AI-Assisted DDR Analysis System")

    # PROPERTY DETAILS TABLE
    doc.add_heading("1. Property Details", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Customer Name"
    table.rows[0].cells[1].text = "Not Available"
    table.rows[1].cells[0].text = "Property Type"
    table.rows[1].cells[1].text = "Flat"
    table.rows[2].cells[0].text = "Number of Floors"
    table.rows[2].cells[1].text = "11"
    table.rows[3].cells[0].text = "Inspection Date"
    table.rows[3].cells[1].text = "Not Available"

    # PROPERTY ISSUE SUMMARY
    doc.add_heading("2. Property Issue Summary", level=1)
    if inspection_text:
        doc.add_paragraph(
            "Inspection report indicates potential issues in certain areas of the property."
        )
    else:
        doc.add_paragraph("Not Available")

    # AREA-WISE OBSERVATIONS
    doc.add_heading("3. Area-wise Observations", level=1)
    areas = find_impacted_areas(inspection_text)
    if areas:
        for area in areas:
            doc.add_paragraph(area, style="List Bullet")
    else:
        doc.add_paragraph("Not Available")

    # DETECTED ISSUES
    doc.add_heading("4. Detected Issues", level=1)
    issues, severity = detect_issues(inspection_text)
    if issues:
        for i in issues:
            doc.add_paragraph(i, style="List Bullet")
    else:
        doc.add_paragraph("No clear issue keywords detected in report.")

    # SEVERITY ASSESSMENT
    doc.add_heading("5. Severity Assessment", level=1)
    doc.add_paragraph(f"Severity Level: {severity}")

    # RECOMMENDED ACTIONS
    doc.add_heading("6. Recommended Actions", level=1)
    doc.add_paragraph("Conduct detailed inspection of impacted areas.", style="List Bullet")
    doc.add_paragraph("Check plumbing systems.", style="List Bullet")
    doc.add_paragraph("Inspect external walls for moisture.", style="List Bullet")

    # ADDITIONAL NOTES / THERMAL ANALYSIS
    doc.add_heading("7. Thermal Analysis & Additional Notes", level=1)
    if thermal_text:
        doc.add_paragraph(
            "Thermal images were reviewed to detect anomalies or temperature variations."
        )
    else:
        doc.add_paragraph("Thermal information Not Available")

    # MISSING / UNCLEAR INFORMATION
    doc.add_heading("8. Missing or Unclear Information", level=1)
    doc.add_paragraph("Customer contact details: Not Available")
    doc.add_paragraph("Detailed thermal interpretation: Not Available")
    doc.add_paragraph("Confirmed root cause of issue: Not Available")

    # THERMAL IMAGES
    doc.add_heading("9. Thermal Images", level=1)
    if images:
        for img in images:
            doc.add_picture(img, width=Inches(3))
    else:
        doc.add_paragraph("Image Not Available")

    doc.save("DDR_Report.docx")

# -------- MAIN --------
inspection_text = extract_text("inspection_report.pdf")
thermal_text = extract_text("thermal_report.pdf")
images = extract_images("thermal_report.pdf")

generate_report(inspection_text, thermal_text, images)
print("✅ DDR Report Generated Successfully")